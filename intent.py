from docx import Document
import schluesselregex as rex
import spacy
from spacy import displacy
from spacy.tokens import Span
from markupsafe import Markup
# from spacy.matcher import PhraseMatcher


import pathlib
import os
import json
import asyncio
import warnings
# from win32com import client
import random
from typing import Dict, Any, List, Tuple


warnings.filterwarnings("ignore", category=SyntaxWarning)
nlp = None

# all_stopwords = nlp.Defaults.stop_words

def spacy_nlp(x):
    global nlp
    return nlp(x)

def spacytest(s: str):
    s1 = remove_stopwords(s)
    print(s1)
    return {"nostop": s1}

def remove_stopwords(word: str) -> str:
    word = word.replace("(", " ")
    word = word.replace(")", " ")
    word = word.replace("/", " ")
    word = word.replace("II", " ")
    word = word.replace("I", " ")
    word = word.replace("Berliner ", " ")
    word = word.replace("GmbH", "")

    wl = spacy_nlp(word)
    tokens = [word for word in wl if not word.is_stop]
    return " ".join(str(x) for x in tokens)

def prepareWords(wordsjs: Dict[str,List[str]]) -> Tuple[Dict[str, Dict[str, str]], Dict[str,List[str]]]:
    # for m in vorhabeninv['intents']:
    #     if (m == "Entwässerung"):
    #         continue
    #     if (m == "Massnahme"):
    #         continue
    #     if (m == "Brandschutz"):
    #         continue
    #     if (m == "Öffentliche Bauten"):
    #         continue
    #     if (m == "Diverse"):
    #         continue
    #     if (m == "Instandsetzung"):
    #         continue
    #     if (m == "Sicherung"):
    #         continue
    #     if (m == "Umbau"):
    #         continue
    #     if (m == "Wasserversorgungs- und Entwässerungsanlagen"):
    #         continue
    #     if (m == "hinterleuchtete , verglaste Großflächenvitrine"):
    #         continue
    #     if (m == "Umbau"):
    #         continue
    #     if (m == "wasserbauliche Anlage"):
    #         continue
    #     words0 = ''
    #     for w in vorhabeninv['intents'][m]:
    #         words0 = words0 + w + ' '
    #     words0 = remove_stopwords(words0)
    #     words0 = words0.replace("- ", " ")
    #     wdoc = spacy_nlp(words0)
    # topics[m] = wdoc

    words: Dict[str, Dict[str, str]] = {}
    wordlist: Dict[str,List[str]] = {}
    for m in wordsjs:
        hierachy: List[str] = wordsjs[m]
        wordlist[m] = hierachy
        dim: str = ""
        if len(hierachy) == 0:
            # dimensions[m] = []
            dim = m
        else:
            dim = hierachy[len(hierachy)-1]
        if (m == "Zimmer"):
            continue
        if (m == "wasserbauliche Anlage"):
            continue
        if (m == "Ort der Leistung"):
            continue
        if (m == "solarthermische Anlage"):
            continue
        if (m == "Maßnahme"):
            continue

           # if (m == "Standsicherheit des Schornsteinkopfes"):
        #     continue
        # if (m == "Stätte der Leistung"):
        #     continue
        # if (m == "Erinnerungs- und Informationsstele"):
        #     continue

        # wdoc = spacy_nlp(m1)
        words[m] = { "dimension": dim}
    return words, wordlist

def similarity(words, wd, wl):
    wdoc: Any = None
    if wd in words:
        w = words[wd]
        if "wdoc" in w:
            wdoc = w["wdoc"]
        else:
            m1 = remove_stopwords(wd)
            m1 = m1.replace("- ", " ")
            m1 = m1.replace(" , ", " ")
            if m1 == " " or len(m1)==0:
                return 0
            wdoc = spacy_nlp(m1)
            w["wdoc"] = wdoc
    if wdoc != None:
        if not wdoc.has_vector:
            print("No vector:", wdoc)
        return wdoc.similarity(wl)
    else:
        return 0

def preparePattern(patternjs) -> List[Dict[str,str]]:
    plist: List[Dict[str,str]] = []
    for p in patternjs:
        patstr: str = p["paragraph"]
        head: str = ""
        tail: str = ""
        pos: int = patstr.find("XXX")
        if pos == 0:
            tail = patstr[4:]
            pos = tail.find("YYY")
            if pos > -1:
                tail = tail[pos+4:]
            plist.append({"head": head, "tail": tail})
        elif pos > -1:
            head = patstr[:pos-1]
            tail = patstr[pos+4:]
            pos = tail.find("YYY")
            if pos > -1:
                tail = tail[pos+4:]
            plist.append({"head": head, "tail": tail})
    return plist

def matchPattern(s, pattern):
    s0 = s
    sl = len(s)
    for p in pattern:
        h = p["head"]
        t = p["tail"]
        dt = s.find(t)
        if dt > 0 and dt == sl-len(t):
            if len(h) > 0:
                if s.find(h) == 0:
                    print("Pattern: ", s[len(h):sl-len(t)])
                    return s[len(h):sl-len(t)]
            else:
                print("Pattern: ", s[:sl-len(t)])
                return s[:sl-len(t)]
    return s0

def extractTopicsAndPlaces(words: Dict[str, Dict[str, str]], wljs, categories, pljs, bljs, bparagraphs: bool, text: str):
    # wordlistjs = wljs
    # pattern = pljs
    # badlist = bljs
    global nlp
    if nlp == None:
        nlp = spacy.load("de_core_news_md")
        # nlp = spacy.load("de")
        nlp.add_pipe('sentencizer')
    if len(text) == 0:
        # try:
            source_dir = r"C:\Data\test\KIbarDok\\txt"
            os.chdir(source_dir)
            files = sorted(os.listdir(source_dir))
            res, all_matches = asyncio.run(extractTopicsAndPlaces1a(files, "Vorhaben:", "Grundstück:", "Grundstücke:", words, wljs,categories, pljs,  bljs, bparagraphs))
            print(all_matches)
            return res
        # except:
        #     pass
    else:
        res, all_matches = asyncio.run(extractTopicsAndPlaces2(
            [], "Vorhaben:", "Grundstück:", "Grundstücke:", words, wljs, categories,  pljs, bljs, text))
        print(all_matches)


         
        # if not nlp.has_pipe("entity_ruler"):
        #     ruler = nlp.add_pipe("entity_ruler")
        # else:
        #     ruler = nlp.get_pipe("entity_ruler")
        # patterns = [
        #             {"label": "ROLLE", "pattern": "Bundeskanzlerin"},
        #             {"label": "ROLLE", "pattern": "Gesundheitsminister"},
        #             # {"label": "GELD", "pattern": "Corona-Bonus"},
        #             # {"label": "GELD", "pattern": [{"LOWER": "Corona"}, {"LOWER": "Bonus"}]},
        #             {"label": "GELD", "pattern": "1000-Euro-Corona-Bonus"}
        #             ]
        # ruler.add_patterns(patterns)
        # ruler.overwrite_ents = False

        doc = nlp(text)

        # new_ents = [x for x in doc.ents]
        # new_ent = Span(doc, 55, 56, label="ACHTUNG")
        # new_ents.append(new_ent)
        
        

        # matcher = PhraseMatcher(nlp.vocab)
        # patterns = [nlp.make_doc(name) for name in ["Pflegenotstand", "Pressekonferenz"]]
        # matcher.add("Names", patterns)
        # for match_id, start, end in matcher(doc):
        #      print("Matched based on lowercase token text:", doc[start:end])
        #      new_ents.append(Span(doc, start, end, label="ISSUE"))

        # doc.ents = new_ents

        # for ent in doc.ents:
        #     print(ent.text, ent.label_)          

        colors = {"ROLLE": "red", "GELD": "yellow"}
        options = {"ents": ["ROLLE", "GELD", "ACHTUNG", "LOC", "ISSUE"], "colors": colors}

        html = displacy.render(doc,style="ent", options=options)
        html = Markup(html.replace("\n\n","\n"))
        return res, html

async def extractTopicsAndPlaces1(files, pattern1, pattern2, pattern2ax, words, wordlistjs, pattern, badlist, bparagraphs):
    tlist = []
    all_matches = {}
    for i in range(0, len(files)):
        ext = files[i][-5:].lower()
        if ext != ".docx":
            continue
        document = Document(files[i])
        topic = ""
        t0 = ""
        wordlist_list = []
        intents = []
        wordlist_list3 = {}
        for p in document.paragraphs:
            pt = p.text
            if len(pt) > 0:
                if pt in badlist:
                    print("Badlist:", pt)
                    continue
                wnfd = False
                pt2 = matchPattern(pt, pattern)
                if pt2 != pt:
                    wnfd = True
                    pt = pt2
                if pt.find("Bearbeiter/in\t\tZimmer") > -1:
                    continue
                if pt.find("Bearbeiter/in\tZimmer") > -1:
                    continue
                if pt.find("Dienstgebäude:\nZimmer") > -1:
                    continue
                if pt.find("Grundstück:") > -1:
                    continue
                if pt.find("Grundstücke:") > -1:
                    continue
                if pt.find("Anlage:") > -1:
                    continue
                if pt.find("Anlagen:") > -1:
                    continue
                pt = pt.replace(" Anlage ", "")

                docp = spacy_nlp(pt)
                # for ent in docp.ents:
                #   for h2 in hida:
                #         h2doc = hida[h2]["nlp"]
                #         h2si = h2doc.similarity(ent)
                #         if h2si > 0.95:
                # print("Hida: ", h2, " in ", ent.lemma_, str(h2si))
                wordlist_list2 = []
                for wl in docp.noun_chunks:
                    w = wl.lemma_
                    w = w.replace("- ", " ")
                    w = w.replace(" , ", " ")
                    w = remove_stopwords(w)
                    # tokens_without_sw= [word for word in text_tokens if not word in all_stopwords]
                    # print(tokens_without_sw)
                    # w = w.replace("die ", " ")
                    # w = w.replace("des ", " ")
                    # w = w.replace("und ", " ")
                    if w.find("Bezirksamt Treptow-Köpenick") > -1:
                        continue
                    if w.find("Am Treptower Park") > -1:
                        continue
                    if w.find("Grundstück") > -1:
                        continue
                    if w.find("Ort") > -1:
                        continue
                    fnd = w in wordlistjs
                    if w in all_matches:
                        w = all_matches[w]["w2"]
                        fnd = True
                    if not fnd and wl.has_vector:
                        matches = {}
                        for w2 in words:
                            if not fnd:
                                # w2doc = words[w2]["wdoc"]
                                # w2si = w2doc.similarity(wl)
                                w2si = similarity(words, w2, wl)
                                # print(w, " ", w2, " ", str(w2si))
                                if w2si > 0.8:
                                    matches[w2si] = w2
                        if len(matches) > 0:
                            w2stlist = sorted(matches, reverse=True)
                            w2si = w2stlist[0]
                            w2 = matches[w2si]
                            all_matches[w] = {"w2": w2, "s": w2si}
                            print(w, " -> ", w2, " (", str(w2si), ")")
                            w = w2
                            fnd = True
                    if fnd:
                        wnfd = True
                        dim = ""
                        if w in words:
                            dim = words[w]["dimension"]
                        if len(dim) > 0:
                            if dim in wordlist_list3:
                                dwl = wordlist_list3[dim]
                                if not w in dwl:
                                    dwl.append(w)
                                    wordlist_list3[dim] = dwl
                            else:
                                wordlist_list3[dim] = [w]
                            if not w in wordlist_list2:
                                wordlist_list2.append(w)
                            if not w in wordlist_list:
                                wordlist_list.append(w)
                        if w in wordlistjs:
                            pl = wordlistjs[w]
                            for pp in pl:
                                if not pp in wordlist_list2:
                                    dim2 = ""
                                    if pp in words:
                                        dim2 = words[pp]["dimension"]
                                    if len(dim2) > 0:
                                        if dim2 in wordlist_list3:
                                            dwl = wordlist_list3[dim2]
                                            if not pp in dwl:
                                                dwl.append(pp)
                                                wordlist_list3[dim2] = dwl
                                        else:
                                            wordlist_list3[dim2] = [pp]
                                        if not pp in wordlist_list2:
                                            wordlist_list2.append(pp)
                                        if not pp in wordlist_list:
                                            wordlist_list.append(pp)
                if wnfd == True:
                    t0 = t0 + "\n" + p.text
                    if bparagraphs:
                        intents.append(
                            {'paragraph': p.text, 'words': wordlist_list2})
        ents = []
        nouns = []
        tfile = files[i]
        doc = spacy_nlp(tfile.replace("_", " ").replace(
            ".docx", "").replace(".", " "))
       # print(tfile + ": " + str(intents))
        print(tfile)
        for e in doc.ents:
            ents.append({'lemma': e.lemma_, 'label': e.label_})
        # for e in doc.noun_chunks:
        #     nouns.append({'lemma': e.lemma_, 'label': e.label_})
        place = ""
        fnd = False
        for p in document.paragraphs:
            txt = p.text
            start_topic = txt.find(pattern1)
            if start_topic != -1:
                fnd = True
                topic = txt[start_topic+10:].split('\n')[0]
                topic = topic.replace("\t", "")
                doc2 = spacy_nlp(topic)
                for e in doc2.ents:
                    ents.append({'lemma': e.lemma_, 'label': e.label_})
                for e in doc2.noun_chunks:
                    nouns.append({'lemma': e.lemma_, 'label': e.label_})
            start_place = txt.find(pattern2)
            if start_place != -1:
                place = txt[start_place+12:].split('\n')[0]
                place = rex.getRegex(place).adresseUnvollständig
            else:
                start_place = txt.find(pattern2ax)
                if start_place != -1:
                    place = txt[start_place+13:].split('\n')[0]
                    place = rex.getRegex(place).adresseUnvollständig
        fnd = True
        if fnd:
            t = {'topic': topic, 'file': tfile,  'place': place, 'keywords': wordlist_list3, 'intents': intents,
                 'entities': ents,  'nouns': nouns}
            try:
                json.dumps(t)
                tlist.append(t)
            except:
                pass
    
    with open('C:\\Data\\test\\topics3.json', 'w', encoding='utf-8') as json_file:
                json.dump(tlist, json_file, indent=4, ensure_ascii=False)
    return tlist, all_matches
    
# entities: ( person: #a6e22d, norp: #e00084, facility: #43c6fc, org: #43c6fc, gpe: #fd9720, loc: #fd9720, product: #8e7dff, event: #ffcc00, work_of_art: #ffcc00, language: #ffcc00, date: #2fbbab, time: #2fbbab, percent: #999, money: #999, quantity: #999, ordinal: #999, cardinal: #999 )

def color_generator(number_of_colors):
        color = ["#"+''.join([random.choice('0123456789ABCDEF') for j in range(6)])
                 for i in range(number_of_colors)]
        return color

def displacyText(pt: str, ents: List[Any], options: Dict[ str, Any]) -> Markup:

    # global nlp
    # if nlp == None:
    #     nlp = spacy.load("de_core_news_md")
    # docp = spacy_nlp(pt)
    # docp.ents = ents
    # html = displacy.render(docp,style="ent", options=options)
    inp: Dict[str, any] = { "text": pt, "ents": ents, "title": None}
    html1: str = displacy.render(inp,style="ent", manual=True, options=options)
    html: Markup = Markup(html1.replace("\n\n","\n"))
    return html

async def extractTopicsAndPlaces1a(files, pattern1, pattern2, pattern2ax, words, wordlistjs, categories, pattern, badlist, bparagraphs):
    tlist: List[Dict] = []
    all_matches = {}

    for i in range(0, len(files)):
        ext: str = files[i][-4:].lower()
        if ext != ".txt":
            continue
        document = ""
        with open(files[i]) as f:
            document = f.read()

        topic = ""
        t0 = ""
        wordlist_list_document = []
        intents = []
        docs_paragraph= []
        wordlist_list_category = {}
        doc = spacy_nlp(document)
        paragraphs= split_in_sentences(document)
        for p in paragraphs:
            pt = p
            if len(pt) > 0:
                if pt in badlist:
                    print("Badlist:", pt)
                    continue
                wnfd = False
                pt2 = matchPattern(pt, pattern)
                if pt2 != pt:
                    wnfd = True
                    pt = pt2
                if pt.find("Bearbeiter/in\t\tZimmer") > -1:
                    continue
                if pt.find("Bearbeiter/in\tZimmer") > -1:
                    continue
                if pt.find("Dienstgebäude:\nZimmer") > -1:
                    continue
                if pt.find("Grundstück:") > -1:
                    continue
                if pt.find("Grundstücke:") > -1:
                    continue
                if pt.find("Anlage:") > -1:
                    continue
                if pt.find("Anlagen:") > -1:
                    continue
                pt = pt.replace(" Anlage ", "")

                docp = spacy_nlp(pt)
                # new_ents = [x for x in docp.ents]
                new_ents = []
                # for ent in docp.ents:
                #   for h2 in hida:
                #         h2doc = hida[h2]["nlp"]
                #         h2si = h2doc.similarity(ent)
                #         if h2si > 0.95:
                # print("Hida: ", h2, " in ", ent.lemma_, str(h2si))
                wordlist_list_paragraph = []
                for wl in docp.noun_chunks:
                    w = wl.lemma_
                    w = w.replace("- ", " ")
                    w = w.replace(" , ", " ")
                    w = remove_stopwords(w)
                    if w.find("Bezirksamt Treptow-Köpenick") > -1:
                        continue
                    if w.find("Am Treptower Park") > -1:
                        continue
                    if w.find("Grundstück") > -1:
                        continue
                    if w.find("Ort") > -1:
                        continue
                    fnd = w in wordlistjs
                    if w in all_matches:
                        w = all_matches[w]["w2"]
                        fnd = True
                    if not fnd and wl.has_vector:
                        matches = {}
                        for w2 in words:
                            if not fnd:
                                # w2doc = words[w2]["wdoc"]
                                # w2si = w2doc.similarity(wl)
                                w2si = similarity(words, w2, wl)
                                # print(w, " ", w2, " ", str(w2si))
                                if w2si > 0.8:
                                    matches[w2si] = w2
                        if len(matches) > 0:
                            w2stlist = sorted(matches, reverse=True)
                            w2si = w2stlist[0]
                            w2 = matches[w2si]
                            all_matches[w] = {"w2": w2, "s": w2si}
                            print(w, " -> ", w2, " (", str(w2si), ")")
                            w = w2
                            fnd = True
                    if fnd:
                        wnfd = True
                        dim = ""
                        if w in words:
                            dim = words[w]["dimension"]
                        if len(dim) > 0:
                            if dim in wordlist_list_category:
                                dwl = wordlist_list_category[dim]
                                if not w in dwl:
                                    dwl.append(w)
                                    wordlist_list_category[dim] = dwl
                            else:
                                wordlist_list_category[dim] = [w]
                            if not w in wordlist_list_paragraph:
                                wordlist_list_paragraph.append(w)
                            if not w in wordlist_list_document:
                                wordlist_list_document.append(w)
                            
                            # new_ents.append(Span(doc, wl.start, wl.end, label=dim.upper()))
                            new_ents.append({ "start": wl.start_char,"end": wl.end_char, "label": dim})
    
                        if w in wordlistjs:
                            pl = wordlistjs[w]
                            for pp in pl:
                                if not pp in wordlist_list_paragraph:
                                    dim2 = ""
                                    if pp in words:
                                        dim2 = words[pp]["dimension"]
                                    if len(dim2) > 0:
                                        if dim2 in wordlist_list_category:
                                            dwl = wordlist_list_category[dim2]
                                            if not pp in dwl:
                                                dwl.append(pp)
                                                wordlist_list_category[dim2] = dwl
                                        else:
                                            wordlist_list_category[dim2] = [pp]
                                        if not pp in wordlist_list_paragraph:
                                            wordlist_list_paragraph.append(pp)
                                        if not pp in wordlist_list_document:
                                            wordlist_list_document.append(pp)
                                        # new_ents.append(Span(doc, wl.start, wl.end, label=w.upper()))
                if wnfd == True:
                    t0 = t0 + "\n" + p
                    # docp.ents = new_ents
                    if bparagraphs:
                        # html = displacy.render(doc,style="ent", options=options)
                        # html = displacy.render(docp,style="ent")
                        # html = Markup(html.replace("\n\n","\n"))
                        intents.append(
                            {'paragraph': p, 'words': wordlist_list_paragraph, "entities": new_ents})
                docs_paragraph.append(docp)
        ents = []
        nouns = []
        tfile = files[i]
        # doc = spacy_nlp(tfile.replace("_", " ").replace(
        #     ".docx", "").replace(".", " "))
       # print(tfile + ": " + str(intents))
        print(tfile)
        for e in doc.ents:
            ents.append({'lemma': e.lemma_, 'label': e.label_})
        # for e in doc.noun_chunks:
        #     nouns.append({'lemma': e.lemma_, 'label': e.label_})
        place = ""
        fnd = False
        for p in paragraphs:
            txt = p
            start_topic = txt.find(pattern1)
            if start_topic != -1:
                fnd = True
                topic = txt[start_topic+10:].split('\n')[0]
                topic = topic.replace("\t", "")
                doc2 = spacy_nlp(topic)
                for e in doc2.ents:
                    ents.append({'lemma': e.lemma_, 'label': e.label_})
                for e in doc2.noun_chunks:
                    nouns.append({'lemma': e.lemma_, 'label': e.label_})
            start_place = txt.find(pattern2)
            if start_place != -1:
                place = txt[start_place+12:].split('\n')[0]
                place = rex.getRegex(place).adresseUnvollständig
            else:
                start_place = txt.find(pattern2ax)
                if start_place != -1:
                    place = txt[start_place+13:].split('\n')[0]
                    place = rex.getRegex(place).adresseUnvollständig
        fnd = True
        if fnd:
            # html = displacy.render(docs_paragraph,style="ent", options=options)
            # html = Markup(html.replace("\n\n","\n"))
            t = {'topic': topic, 'file': tfile.replace(".txt",".docx"),  'place': place, 
                 'keywords': wordlist_list_category, 
                 'intents': intents,
                'nouns': nouns }
 # 'entities': ents,
            try:
                json.dumps(t)
                tlist.append(t)
            except:
                pass
    
    with open('C:\\Data\\test\\topics3a.json', 'w', encoding='utf-8') as json_file:
                json.dump(tlist, json_file, indent=4, ensure_ascii=False)
    return tlist, all_matches

def split_in_sentences(text):
    doc = spacy_nlp(text)
    return [str(sent).strip() for sent in doc.sents]

async def extractTopicsAndPlaces2(files, pattern1, pattern2, pattern2ax, words, wordlistjs, pattern, badlist, text):
    tlist = []
    all_matches = {}
    topic = ""
    t0 = ""
    wordlist_list = []
    intents = []
    wordlist_list3 = {}
    paragraphs = split_in_sentences(text)

    for pt in paragraphs:
        if len(pt) > 0:
            if pt in badlist:
                print("Badlist:", pt)
                continue
            pt2 = matchPattern(pt, pattern)
            if pt2 != pt:
                wnfd = True
                pt = pt2
            if pt.find("Bearbeiter/in\t\tZimmer") > -1:
                continue
            if pt.find("Bearbeiter/in\tZimmer") > -1:
                continue
            if pt.find("Dienstgebäude:\nZimmer") > -1:
                continue
            if pt.find("Grundstück:") > -1:
                continue
            if pt.find("Grundstücke:") > -1:
                continue
            if pt.find("Anlage:") > -1:
                continue
            if pt.find("Anlagen:") > -1:
                continue
            pt = pt.replace(" Anlage ", "")

            docp = spacy_nlp(pt)
            # for ent in docp.ents:
            #   for h2 in hida:
            #         h2doc = hida[h2]["nlp"]
            #         h2si = h2doc.similarity(ent)
            #         if h2si > 0.95:
            # print("Hida: ", h2, " in ", ent.lemma_, str(h2si))
            wnfd = False
            wordlist_list2 = []
            for wl in docp.noun_chunks:
                w = wl.lemma_
                w = w.replace("- ", " ")
                w = w.replace(" , ", " ")
                w = remove_stopwords(w)
                # tokens_without_sw= [word for word in text_tokens if not word in all_stopwords]
                # print(tokens_without_sw)
                # w = w.replace("die ", " ")
                # w = w.replace("des ", " ")
                # w = w.replace("und ", " ")
                if w.find("Bezirksamt Treptow-Köpenick") > -1:
                    continue
                if w.find("Am Treptower Park") > -1:
                    continue
                if w.find("Grundstück") > -1:
                    continue
                if w.find("Ort") > -1:
                    continue
                fnd = w in wordlistjs
                if w in all_matches:
                    w = all_matches[w]["w2"]
                    fnd = True
                if not fnd  and wl.has_vector:
                    matches = {}
                    for w2 in words:
                        if not fnd:
                            # w2doc = words[w2]["wdoc"]
                            # w2si = w2doc.similarity(wl)
                            w2si = similarity(words, w2, wl)
                            # print(w, " ", w2, " ", str(w2si))
                            if w2si > 0.8:
                                matches[w2si] = w2
                    if len(matches) > 0:
                        w2stlist = sorted(matches, reverse=True)
                        w2si = w2stlist[0]
                        w2 = matches[w2si]
                        all_matches[w] = {"w2": w2, "s": w2si}
                        print(w, " -> ", w2, " (", str(w2si), ")")
                        w = w2
                        fnd = True
                if fnd:
                    wnfd = True
                    dim = ""
                    if w in words:
                        dim = words[w]["dimension"]
                    if len(dim) > 0:
                        if dim in wordlist_list3:
                            dwl = wordlist_list3[dim]
                            if not w in dwl:
                                dwl.append(w)
                                wordlist_list3[dim] = dwl
                        else:
                            wordlist_list3[dim] = [w]
                        if not w in wordlist_list2:
                            wordlist_list2.append(w)
                        if not w in wordlist_list:
                            wordlist_list.append(w)
                    if w in wordlistjs:
                        pl = wordlistjs[w]
                        for pp in pl:
                            if not pp in wordlist_list2:
                                dim2 = ""
                                if pp in words:
                                    dim2 = words[pp]["dimension"]
                                if len(dim2) > 0:
                                    if dim2 in wordlist_list3:
                                        dwl = wordlist_list3[dim2]
                                        if not pp in dwl:
                                            dwl.append(pp)
                                            wordlist_list3[dim2] = dwl
                                    else:
                                        wordlist_list3[dim2] = [pp]
                                    if not pp in wordlist_list2:
                                        wordlist_list2.append(pp)
                                    if not pp in wordlist_list:
                                        wordlist_list.append(pp)
            if wnfd == True:
                t0 = t0 + "\n" + pt
                intents.append(
                    {'paragraph': pt, 'words': wordlist_list2})
        ents = []
        nouns = []
    #     tfile = files[i]
    #     doc = spacy_nlp(tfile.replace("_", " ").replace(
    #         ".docx", "").replace(".", " "))
    #    # print(tfile + ": " + str(intents))
    # print(tfile)
    # for e in doc.ents:
    #     ents.append({'lemma': e.lemma_, 'label': e.label_})
    # for e in doc.noun_chunks:
    #     nouns.append({'lemma': e.lemma_, 'label': e.label_})
    place = ""
    fnd = False
    for txt in paragraphs:
        start_topic = txt.find(pattern1)
        if start_topic != -1:
            fnd = True
            topic = txt[start_topic+10:].split('\n')[0]
            topic = topic.replace("\t", "")
            doc2 = spacy_nlp(topic)
            for e in doc2.ents:
                ents.append({'lemma': e.lemma_, 'label': e.label_})
            for e in doc2.noun_chunks:
                nouns.append({'lemma': e.lemma_, 'label': e.label_})
        start_place = txt.find(pattern2)
        if start_place != -1:
            place = txt[start_place+12:].split('\n')[0]
            place = rex.getRegex(place).adresseUnvollständig
        else:
            start_place = txt.find(pattern2ax)
            if start_place != -1:
                place = txt[start_place+13:].split('\n')[0]
                place = rex.getRegex(place).adresseUnvollständig
    fnd = True
    if fnd:
        t = {'topic': topic, 'file': "",  'place': place, 'keywords': wordlist_list3, 'intents': intents,
             'entities': ents,  'nouns': nouns, 'text': text}
        try:
            json.dumps(t)
            tlist.append(t)
        except:
            pass
    return tlist, all_matches

def extractText():
        # try:
            target_dir = r"C:\Data\test\KIbarDok\\docx"
            os.chdir(target_dir)
            files = sorted(os.listdir(target_dir))
            res = asyncio.run(extractText1(files))
            return res
        # except:
        #     pass

async def extractText1(files):
    tlist = []
    for i in range(0, len(files)):
        ext = files[i][-5:].lower()
        if ext != ".docx":
            continue
        try:
            document = Document(files[i])
        except:
            continue
        doctext = ""
        for p in document.paragraphs:
            txt = p.text
            if len(txt) > 0:
                doctext = doctext + "\n" + txt
        if len(doctext)>0:
            doctext= doctext[1:]
        tfile = files[i]
        print(tfile)
        t = {'file': tfile, 'text': doctext}
        try:
            json.dumps(t)
            tlist.append(t)
        except:
            pass
    with open('C:\\Data\\test\\text3.json', 'w', encoding='utf-8') as json_file:
                json.dump(tlist, json_file, indent=4, ensure_ascii=False)
    return tlist
 
#  extractText();

def doc2docx(word, doc_path, docx_path):
    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, 16)  # 16 doc2docx
        doc.Close()
    except:
        pass

#  https://docs.microsoft.com/en-us/previous-versions/office/developer/office-2003/aa220734(v=office.11)
#  https://docs.microsoft.com/en-us/office/vba/api/word.wdsaveformat
def doc2txt(word, doc_path, docx_path):
    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, 2)  # wdFormatText
        doc.Close()
    except:
        pass
 
def doc2pdf(word, doc_path, docx_path):
    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(docx_path, 17)  # pdf
        doc.Close()
    except:
        pass

# def convert():
#     word = client.Dispatch('word.Application')

#     # path = "C:\\Data\\test\\KIbarDok\\docx"
#     # docx_path = 'C:\\Data\\test\\KIbarDok\\pdf'

#     # i = 0
#     # for root, d_names, f_names in os.walk(path):
#     #     for f in f_names:
#     #         if f.endswith(".docx"):
#     #             i = i+1
#     #             if i > 770:
#     #                 print(i, " ", f)
#     #                 doc2pdf(word, os.path.join(root, f), os.path.join(
#     #                     docx_path, f.replace(".docx", ".pdf")))  

#     path = "C:\\Data\\test\\KIbarDok\\docx"
#     docx_path = 'C:\\Data\\test\\KIbarDok\\txt'

#     i = 0
#     for root, d_names, f_names in os.walk(path):
#         for f in f_names:
#             if f.endswith(".docx"):
#                 i = i+1
#                 if i > 730:
#                     print(i, " ", f)
#                     doc2txt(word, os.path.join(root, f), os.path.join(
#                         docx_path, f.replace(".docx", ".txt")))  

#     word.Quit()
